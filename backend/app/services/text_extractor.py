from PyPDF2 import PdfReader
from typing import Optional
import io

async def extract_pdf_text(file_content: bytes) -> dict:
    """
    Extract text from PDF file
    Returns: {text, page_count, metadata}
    """
    try:
        # Create PDF reader from bytes
        pdf_file = io.BytesIO(file_content)
        reader = PdfReader(pdf_file)
        
        # Get metadata
        page_count = len(reader.pages)
        metadata = reader.metadata if reader.metadata else {}
        
        # Extract text from all pages
        text_parts = []
        for page in reader.pages:
            text = page.extract_text()
            if text:
                text_parts.append(text)
        
        full_text = "\n\n".join(text_parts)
        
        return {
            "text": full_text,
            "page_count": page_count,
            "metadata": {
                "author": metadata.get("/Author", "Unknown"),
                "title": metadata.get("/Title", ""),
                "subject": metadata.get("/Subject", ""),
                "creator": metadata.get("/Creator", ""),
            },
            "success": True,
            "method": "pdf"
        }
    
    except Exception as e:
        return {
            "text": "",
            "page_count": 0,
            "metadata": {},
            "success": False,
            "error": str(e),
            "method": "pdf"
        }

async def extract_image_text(file_content: bytes) -> dict:
    """
    Extract text from image using OCR (Tesseract)
    """
    try:
        from PIL import Image
        import pytesseract
        import io
        
        # Open image from bytes
        image = Image.open(io.BytesIO(file_content))
        
        # Perform OCR
        text = pytesseract.image_to_string(image)
        
        # Get confidence data
        data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
        confidences = [int(conf) for conf in data['conf'] if conf != '-1']
        avg_confidence = sum(confidences) / len(confidences) if confidences else 0
        
        return {
            "text": text.strip(),
            "success": True,
            "method": "ocr",
            "confidence": avg_confidence,
            "language": "eng"  # Default to English
        }
    
    except Exception as e:
        print(f"OCR Error: {e}")
        return {
            "text": "",
            "success": False,
            "method": "ocr",
            "error": str(e)
        }

async def extract_docx_text(file_content: bytes) -> dict:
    """
    Extract text from DOCX file
    """
    try:
        from docx import Document
        import io
        
        # Load document from bytes
        doc = Document(io.BytesIO(file_content))
        
        # Extract text from paragraphs
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        full_text = "\n\n".join(paragraphs)
        
        # Extract text from tables
        tables_text = []
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text for cell in row.cells]
                tables_text.append(" | ".join(row_text))
        
        if tables_text:
            full_text += "\n\n" + "\n".join(tables_text)
        
        return {
            "text": full_text.strip(),
            "success": True,
            "method": "docx",
            "paragraph_count": len(paragraphs),
            "table_count": len(doc.tables)
        }
    
    except Exception as e:
        print(f"DOCX Error: {e}")
        return {
            "text": "",
            "success": False,
            "method": "docx",
            "error": str(e)
        }

async def extract_text(file_content: bytes, file_type: str) -> Optional[str]:
    """
    Extract text based on file type
    """
    file_type_lower = file_type.lower()
    
    if 'pdf' in file_type_lower:
        result = await extract_pdf_text(file_content)
        return result.get("text", "") if result.get("success") else None
    
    elif any(img_type in file_type_lower for img_type in ['image', 'jpeg', 'jpg', 'png']):
        result = await extract_image_text(file_content)
        return result.get("text", "") if result.get("success") else None
    
    elif 'word' in file_type_lower or 'document' in file_type_lower or file_type_lower.endswith('.docx'):
        result = await extract_docx_text(file_content)
        return result.get("text", "") if result.get("success") else None
    
    return None
